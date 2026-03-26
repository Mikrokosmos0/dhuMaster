import os
import sys
import json
from collections import Counter, defaultdict


def ensure_pydeps_on_path():
    """
    将本仓库下的 .pydeps 目录加入 sys.path，便于导入 python-docx。
    """
    root = os.path.dirname(os.path.abspath(__file__))
    pydeps = os.path.join(root, ".pydeps")
    if os.path.isdir(pydeps) and pydeps not in sys.path:
        sys.path.insert(0, pydeps)


def load_document(docx_path):
    from docx import Document

    if not os.path.exists(docx_path):
        raise FileNotFoundError(f"Docx 文件不存在: {docx_path}")
    return Document(docx_path)


def length_to_pt(length):
    """
    将 python-docx 的长度对象或数值近似转换为 pt，返回 float 或 None。
    """
    if length is None:
        return None
    try:
        # 部分字段本身就是 float（如多倍行距）
        if isinstance(length, (int, float)):
            return float(length)
        # python-docx 的 Length 有 .pt 属性
        pt = getattr(length, "pt", None)
        if pt is not None:
            return float(pt)
    except Exception:
        pass
    return None


def length_to_cm(length):
    """
    将 python-docx 的长度对象近似转换为 cm，返回 float 或 None。
    """
    if length is None:
        return None
    try:
        cm = getattr(length, "cm", None)
        if cm is not None:
            return float(cm)
    except Exception:
        pass
    # 兜底：如果有 .pt，则用 1pt ≈ 0.0352778cm
    pt = length_to_pt(length)
    if pt is not None:
        return pt * 0.0352778
    return None


def extract_paragraph_info(doc):
    paragraphs_info = []

    for idx, p in enumerate(doc.paragraphs):
        text = p.text or ""
        style_name = p.style.name if p.style is not None else None

        # 统计本段中各 run 的字体和字号
        font_names = []
        font_sizes = []
        bold_flags = []
        for r in p.runs:
            if r.font is not None:
                if r.font.name:
                    font_names.append(r.font.name)
                if r.font.size:
                    font_sizes.append(length_to_pt(r.font.size))
                bold_flags.append(bool(r.font.bold))

        def most_common_or_none(values):
            values = [v for v in values if v is not None]
            if not values:
                return None
            counter = Counter(values)
            return counter.most_common(1)[0][0]

        main_font = most_common_or_none(font_names)
        main_font_size = most_common_or_none(font_sizes)
        is_bold = None
        if bold_flags:
            # 视 True 比 False 权重大
            true_count = sum(1 for b in bold_flags if b)
            false_count = len(bold_flags) - true_count
            is_bold = true_count >= false_count

        pf = p.paragraph_format
        align = getattr(p.alignment, "name", None) if p.alignment is not None else None
        line_spacing = pf.line_spacing
        line_spacing_rule = getattr(pf.line_spacing_rule, "name", None) if pf.line_spacing_rule is not None else None
        space_before = length_to_pt(pf.space_before)
        space_after = length_to_pt(pf.space_after)

        # 简单根据样式名推测是否为标题/级别
        is_heading = False
        heading_level = None
        heading_keywords = ["Heading", "标题", "Title"]
        if style_name:
            if any(k in style_name for k in heading_keywords):
                is_heading = True
                # 尝试从样式名中提取级别数字
                for ch in style_name:
                    if ch.isdigit():
                        heading_level = int(ch)
                        break
                if "Title" in style_name or "标题" in style_name and heading_level is None:
                    heading_level = 0

        paragraphs_info.append(
            {
                "index": idx,
                "text": text,
                "style_name": style_name,
                "is_heading": is_heading,
                "heading_level": heading_level,
                "font_name": main_font,
                "font_size_pt": main_font_size,
                "bold": is_bold,
                "alignment": align,
                "line_spacing": line_spacing if isinstance(line_spacing, (int, float)) else length_to_pt(line_spacing),
                "line_spacing_rule": line_spacing_rule,
                "space_before_pt": space_before,
                "space_after_pt": space_after,
            }
        )

    return paragraphs_info


def extract_table_info(doc):
    tables_info = []
    for t_idx, tbl in enumerate(doc.tables):
        style_name = tbl.style.name if tbl.style is not None else None
        rows_data = []
        for r in tbl.rows:
            row_cells = []
            for c in r.cells:
                row_cells.append(c.text or "")
            rows_data.append(row_cells)
        tables_info.append(
            {
                "index": t_idx,
                "style_name": style_name,
                "rows": rows_data,
            }
        )
    return tables_info


def extract_section_page_setup(doc):
    """
    提取文档节的页面设置（以第一节为主）。
    """
    if not doc.sections:
        return {}
    sec = doc.sections[0]
    return {
        "page_width_cm": length_to_cm(sec.page_width),
        "page_height_cm": length_to_cm(sec.page_height),
        "top_margin_cm": length_to_cm(sec.top_margin),
        "bottom_margin_cm": length_to_cm(sec.bottom_margin),
        "left_margin_cm": length_to_cm(sec.left_margin),
        "right_margin_cm": length_to_cm(sec.right_margin),
        "gutter_cm": length_to_cm(sec.gutter),
        "orientation": getattr(sec.orientation, "name", None) if getattr(sec, "orientation", None) is not None else None,
    }


def build_styles_summary(paragraphs):
    """
    汇总段落中使用到的样式名称及其计数。
    """
    style_counter = Counter()
    for p in paragraphs:
        if p["style_name"]:
            style_counter[p["style_name"]] += 1
    return {
        "paragraph_styles": [{"name": name, "count": count} for name, count in style_counter.most_common()]
    }


def write_clean_text(paragraphs, tables, output_txt_path):
    """
    输出较为干净的文本视图：段落逐行输出，表格以简单文本形式插入。
    """
    lines = []
    para_idx_to_text = {p["index"]: p["text"] for p in paragraphs}

    # 直接按段落顺序输出
    for idx in sorted(para_idx_to_text.keys()):
        text = para_idx_to_text[idx]
        lines.append(text)

    # 之后附加表格内容，便于人工查看
    if tables:
        lines.append("")
        lines.append("========== TABLES (文本视图，仅供参考) ==========")
        for tbl in tables:
            lines.append(f"[表格 {tbl['index']}, 样式: {tbl['style_name']}]")
            for row in tbl["rows"]:
                lines.append(" | ".join(row))
            lines.append("")

    with open(output_txt_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))


def generate_format_spec_from_json(extracted_json_path, format_spec_path):
    """
    根据 requirements_extracted.json 自动分析并生成格式规范描述 JSON。
    该规范会尽可能依据实际使用的样式和字号等信息，不引入任何外部假设。
    """
    with open(extracted_json_path, "r", encoding="utf-8") as f:
        data = json.load(f)

    paragraphs = data.get("paragraphs", [])
    page_setup = data.get("page_setup", {})

    # 区分标题段落和正文段落
    heading_paras = [p for p in paragraphs if p.get("is_heading")]
    body_paras = [p for p in paragraphs if not p.get("is_heading") and (p.get("text") or "").strip()]

    def aggregate_numeric(values):
        vals = [v for v in values if isinstance(v, (int, float))]
        if not vals:
            return None
        return sum(vals) / len(vals)

    def most_common(values):
        vals = [v for v in values if v is not None]
        if not vals:
            return None
        c = Counter(vals)
        return c.most_common(1)[0][0]

    # 正文格式统计
    body_font_names = [p.get("font_name") for p in body_paras]
    body_font_sizes = [p.get("font_size_pt") for p in body_paras]
    body_aligns = [p.get("alignment") for p in body_paras]
    body_line_spacings = [p.get("line_spacing") for p in body_paras]
    body_space_before = [p.get("space_before_pt") for p in body_paras]
    body_space_after = [p.get("space_after_pt") for p in body_paras]

    body_spec = {
        "font_name": most_common(body_font_names),
        "font_size_pt": aggregate_numeric(body_font_sizes),
        "alignment": most_common(body_aligns),
        "line_spacing": aggregate_numeric(body_line_spacings),
        "space_before_pt": aggregate_numeric(body_space_before),
        "space_after_pt": aggregate_numeric(body_space_after),
    }

    # 标题分级统计
    headings_by_level = defaultdict(list)
    for p in heading_paras:
        lvl = p.get("heading_level")
        if lvl is None:
            # 视为 0 级
            lvl = 0
        headings_by_level[lvl].append(p)

    headings_spec = {}
    for lvl, items in sorted(headings_by_level.items(), key=lambda kv: kv[0]):
        names = [p.get("font_name") for p in items]
        sizes = [p.get("font_size_pt") for p in items]
        aligns = [p.get("alignment") for p in items]
        line_spacings = [p.get("line_spacing") for p in items]
        bold_flags = [p.get("bold") for p in items]
        space_before = [p.get("space_before_pt") for p in items]
        space_after = [p.get("space_after_pt") for p in items]

        true_bold = sum(1 for b in bold_flags if b)
        false_bold = sum(1 for b in bold_flags if b is False)
        bold_value = None
        if true_bold or false_bold:
            bold_value = true_bold >= false_bold

        headings_spec[str(lvl)] = {
            "font_name": most_common(names),
            "font_size_pt": aggregate_numeric(sizes),
            "alignment": most_common(aligns),
            "line_spacing": aggregate_numeric(line_spacings),
            "bold": bold_value,
            "space_before_pt": aggregate_numeric(space_before),
            "space_after_pt": aggregate_numeric(space_after),
        }

    # 样式使用汇总直接从提取 JSON 拷贝
    styles_summary = data.get("styles_summary", {})

    format_spec = {
        "source": "requirements.docx",
        "page_setup": page_setup,
        "body": body_spec,
        "headings": headings_spec,
        "styles_summary": styles_summary,
    }

    with open(format_spec_path, "w", encoding="utf-8") as f:
        json.dump(format_spec, f, ensure_ascii=False, indent=2)


def main():
    root = os.path.dirname(os.path.abspath(__file__))
    docx_path = os.path.join(root, "requirements.docx")
    txt_out = os.path.join(root, "requirements_extracted.txt")
    json_out = os.path.join(root, "requirements_extracted.json")
    spec_out = os.path.join(root, "format_spec.json")

    ensure_pydeps_on_path()
    doc = load_document(docx_path)

    paragraphs = extract_paragraph_info(doc)
    tables = extract_table_info(doc)
    page_setup = extract_section_page_setup(doc)
    styles_summary = build_styles_summary(paragraphs)

    # 写出结构化 JSON
    extracted = {
        "source": os.path.basename(docx_path),
        "page_setup": page_setup,
        "paragraphs": paragraphs,
        "tables": tables,
        "styles_summary": styles_summary,
    }

    with open(json_out, "w", encoding="utf-8") as f:
        json.dump(extracted, f, ensure_ascii=False, indent=2)

    # 写出纯文本视图
    write_clean_text(paragraphs, tables, txt_out)

    # 基于刚写出的 JSON 再生成格式规范 JSON
    generate_format_spec_from_json(json_out, spec_out)


if __name__ == "__main__":
    main()

